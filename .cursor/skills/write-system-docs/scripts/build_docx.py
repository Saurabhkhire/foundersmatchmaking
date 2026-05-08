from pathlib import Path
import sys

try:
    from docx import Document
except Exception:
    raise SystemExit("python-docx is required. Install with: pip install python-docx")


def markdown_to_docx(md_dir: Path, output_name: str) -> None:
    doc = Document()
    doc.add_heading("System Documentation", level=1)

    files = sorted(md_dir.glob("*.md"))
    if not files:
        raise SystemExit(f"No markdown files found in {md_dir}")

    for md_file in files:
        doc.add_page_break()
        doc.add_heading(md_file.stem.replace("-", " ").title(), level=2)
        for line in md_file.read_text(encoding="utf-8").splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)

    out_path = md_dir / output_name
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: python build_docx.py <markdown_dir> <output_file_name>")
    markdown_to_docx(Path(sys.argv[1]), sys.argv[2])
